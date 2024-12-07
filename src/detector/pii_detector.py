"""
PII Detection Module

This module detects PII patterns in text.
Pattern matching is performed based on pattern definition files using regular expressions.

PII 검출 모듈

이 모듈은 텍스트에서 PII 패턴을 검출합니다.
패턴 정의 파일을 기반으로 정규표현식 매칭을 수행합니다.

Usage example / 사용 예시:
    detector = PIIDetector("patterns.json")
    results = detector.detect("주민번호는 xxxxxxx-xxxxxxxx 입니다")
"""

import re
import magic
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass
from .patterns import PatternLoader, PIIPattern
import logging


logger = logging.getLogger(__name__)

@dataclass
class DetectionResult:
    """
    Data class containing PII detection results.
    This class stores detailed information about each PII pattern match found during scanning.

    Attributes:
        pattern_name: Name of the detected pattern (e.g., "SSN", "Credit Card")
        matched_text: The actual text segment that matched the pattern
        start_pos: Starting character position of the match in the text
        end_pos: Ending character position of the match in the text
        risk_level: Risk level of the detected pattern (HIGH/MEDIUM/LOW)
        line_number: Line number where the pattern was found
        context: Surrounding text context of the matched pattern
        file_path: Path to the source file where the pattern was found (optional)

    PII 검출 결과를 담는 데이터 클래스입니다.
    이 클래스는 스캔 중 발견된 각 PII 패턴 매치에 대한 상세 정보를 저장합니다.

    Attributes:
        pattern_name: 검출된 패턴의 이름 (예: "주민등록번호", "신용카드")
        matched_text: 패턴과 일치한 실제 텍스트 부분
        start_pos: 텍스트에서 매치가 시작되는 문자 위치
        end_pos: 텍스트에서 매치가 끝나는 문자 위치
        risk_level: 검출된 패턴의 위험도 (HIGH/MEDIUM/LOW)
        line_number: 패턴이 발견된 줄 번호
        context: 매치된 패턴 주변의 문맥 텍스트
        file_path: 패턴이 발견된 원본 파일 경로 (선택사항)
    """
    pattern_name: str
    matched_text: str
    start_pos: int
    end_pos: int
    risk_level: str
    line_number: int
    context: str
    file_path: Path = None

class PIIDetector:
    """
    Class for detecting PII patterns
    PII 패턴을 검출하는 클래스
    """
    
    def __init__(self, pattern_file: str):
        """
        Initialize PIIDetector

        Args:
            pattern_file: Path to JSON file containing PII pattern definitions
            
        PIIDetector 초기화

        Args:
            pattern_file: PII 패턴 정의가 포함된 JSON 파일 경로
        """
        self.patterns = PatternLoader.load_patterns(pattern_file)
        self.compiled_patterns = {
            pattern.name: re.compile(pattern.pattern) 
            for pattern in self.patterns
        }
    
    def detect(self, text: str, context_chars: int = 50) -> List[DetectionResult]:
        """
        Detect PII patterns in text

        Args:
            text: Text to scan for PII
            context_chars: Number of characters to include as context around matches

        Returns:
            List of DetectionResult objects

        텍스트에서 PII 패턴을 검출합니다.

        Args:
            text: 검사할 텍스트
            context_chars: 매칭된 텍스트 주변의 문맥을 포함할 문자 수

        Returns:
            DetectionResult 객체들의 리스트
        """
        results = []
        lines = text.splitlines()
        
        for line_num, line in enumerate(lines, 1):
            for pattern in self.patterns:
                matches = self.compiled_patterns[pattern.name].finditer(line)
                for match in matches:
                    start, end = match.span()
                    # 매칭된 부분 주변의 문맥 추출
                    context_start = max(0, start - context_chars)
                    context_end = min(len(line), end + context_chars)
                    context = line[context_start:context_end]
                    
                    result = DetectionResult(
                        pattern_name=pattern.name,
                        matched_text=match.group(),
                        start_pos=start,
                        end_pos=end,
                        risk_level=pattern.risk_level.value,
                        line_number=line_num,
                        context=context
                    )
                    results.append(result)
        
        return results

    def detect_file(self, file_path: str) -> List[DetectionResult]:
        """
        Detect PII patterns in a file.

        Args:
            file_path: Path to the file to scan

        Returns:
            List of DetectionResult objects

        파일에서 PII 패턴을 검출합니다.

        Args:
            file_path: 검사할 파일 경로

        Returns:
            DetectionResult 객체들의 리스트
        """
        file_path = Path(file_path)
        mime_type = magic.from_file(str(file_path), mime=True)
        logger.debug(f"파일 MIME 타입: {mime_type} - {file_path}")
        
        try:
            if 'pdf' in mime_type.lower():
                logger.debug(f"PDF 파일 처리 중: {file_path}")
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
            
            elif 'officedocument.wordprocessingml.document' in mime_type.lower() or file_path.suffix.lower() == '.docx':
                logger.debug(f"DOCX 파일 처리 중: {file_path}")
                import docx
                doc = docx.Document(file_path)
                
                # 일반 텍스트 추출
                paragraphs_text = [paragraph.text for paragraph in doc.paragraphs]
                
                # 테이블 내용 추출
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        tables_text.append(' '.join(row_text))
                
                # 모든 텍스트 결합
                text = '\n'.join(paragraphs_text + tables_text)
                logger.debug(f"DOCX 텍스트 추출 완료: {len(text)} 문자")
            
            else:
                logger.debug(f"일반 텍스트 파일 처리 중: {file_path}")
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            results = self.detect(text)
            for result in results:
                result.file_path = file_path
                
            return results
            
        except Exception as e:
            logger.error(f"파일 처리 중 오류 발생 - {file_path}: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())
            return []
