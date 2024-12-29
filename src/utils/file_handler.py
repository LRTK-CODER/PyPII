from utils.logger import setup_logger
from enum import Enum
from typing import Optional, Union
from pathlib import Path
import magic

logger = setup_logger(__name__)

class FileType(Enum):
    """
    Enumeration of supported file types. / 지원되는 파일 유형의 열거형입니다.
    
    Attributes:
        TEXT: Plain text files / 일반 텍스트 파일
        PDF: PDF documents / PDF 문서
        DOCX: Microsoft Word documents / 마이크로소프트 워드 문서
        UNKNOWN: Unsupported file types / 지원되지 않는 파일 유형
    """
    TEXT = "text"  
    PDF = "PDF"
    DOCX = "docx"
    UNKNOWN = "unknown"

class FileHandler:
    """
    Handles file operations and text extraction from various file formats. / 다양한 파일 형식에서 파일 작업 및 텍스트 추출을 처리합니다.
    """

    @staticmethod
    def get_file_type(file_path: Union[str, Path]) -> FileType:
        """
        Determines the type of the given file based on its MIME type. / 파일의 MIME 타입을 기반으로 파일 유형을 결정합니다.

        Args:
            file_path: Path to the file to check / 확인할 파일의 경로

        Returns:
            FileType: The detected file type / 감지된 파일 유형
        """
        path = Path(file_path)
        mime_type = magic.from_file(str(path), mime=True).lower()
        logger.debug(f"Detected MIME type for {path}: {mime_type}")

        if "text/plain" in mime_type:
            return FileType.TEXT
        elif "pdf" in mime_type:
            return FileType.PDF
        elif "officedocument.wordprocessingml.document" in mime_type or path.suffix.lower() == ".docx":
            return FileType.DOCX
        else:
            logger.warning(f"Unsupported file type detected: {mime_type} for {path}")
            return FileType.UNKNOWN

    @staticmethod
    def _extract_from_text(file_path: Path) -> str:
        """
        Extracts text content from a plain text file. / 일반 텍스트 파일에서 텍스트 내용을 추출합니다.

        Args:
            file_path: Path to the text file / 텍스트 파일의 경로

        Returns:
            str: The extracted text content / 추출된 텍스트 내용
        """
        logger.debug(f"Extracting text from text file: {file_path}")
        return file_path.read_text(encoding="utf-8")
    
    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        """
        Extracts text content from a PDF file using pdfplumber. / pdfplumber를 사용하여 PDF 파일에서 텍스트 내용을 추출합니다.

        Args:
            file_path: Path to the PDF file / PDF 파일의 경로

        Returns:
            str: The extracted text content / 추출된 텍스트 내용

        Raises:
            ImportError: If pdfplumber package is not installed / pdfplumber 패키지가 설치되지 않은 경우
        """
        try:
            import pdfplumber
            logger.debug(f"Extracting text from PDF file: {file_path}")
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(
                    page.extract_text() or ""
                    for page in pdf.pages
                )
                logger.debug(f"Successfully extracted {len(text)} characters from PDF")
                return text
        except ImportError:
            logger.error("pdfplumber package is required for PDF support")
            raise

    @staticmethod
    def _extract_from_docx(file_path: Path) -> str:
        """
        Extracts text content from a DOCX file using python-docx. / python-docx를 사용하여 DOCX 파일에서 텍스트 내용을 추출합니다.

        Args:
            file_path: Path to the DOCX file / DOCX 파일의 경로

        Returns:
            str: The extracted text content / 추출된 텍스트 내용

        Raises:
            ImportError: If python-docx package is not installed / python-docx 패키지가 설치되지 않은 경우
        """
        try:
            import docx
            logger.debug(f"Extracting text from DOCX file: {file_path}")
            doc = docx.Document(file_path)

            # Extract text from paragraphs / 단락에서 텍스트 추출
            paragraphs = [p.text for p in doc.paragraphs]

            # Extract text from tables / 표에서 텍스트 추출
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" ".join(cell.text for cell in row.cells))

            text = "\n".join(paragraphs + tables)
            logger.debug(f"Successfully extracted {len(text)} characters from DOCX")
            return text
        except ImportError:
            logger.error("python-docx package is required for DOCX support")
            raise

    @staticmethod
    def extract_text(file_path: Union[str, Path]) -> Optional[str]:
        """
        Extracts text content from a file based on its type. / 파일 유형에 따라 텍스트 내용을 추출합니다.

        Args:
            file_path: Path to the file to extract text from / 텍스트를 추출할 파일의 경로

        Returns:
            Optional[str]: The extracted text content, or None if extraction fails / 추출된 텍스트 내용, 추출 실패 시 None 반환

        Raises:
            ValueError: If the file type is not supported / 지원되지 않는 파일 유형인 경우
        """
        path = Path(file_path)
        file_type = FileHandler.get_file_type(path)
        logger.info(f"Attempting to extract text from {path} (type: {file_type.value})")

        try:
            if file_type == FileType.TEXT:
                return FileHandler._extract_from_text(path)
            elif file_type == FileType.PDF:
                return FileHandler._extract_from_pdf(path)
            elif file_type == FileType.DOCX:
                return FileHandler._extract_from_docx(path)
            else:
                logger.error(f"Unsupported file type: {path.suffix}")
                raise ValueError(f"Unsupported file type: {path.suffix}")
        except Exception as e:
            logger.error(f"Failed to extract text from {path}: {str(e)}")
            return None