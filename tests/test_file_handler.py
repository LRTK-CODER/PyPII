import pytest
import magic
from pathlib import Path
from src.utils.file_handler import FileHandler, FileType

@pytest.fixture
def sample_text_file(tmp_path):
    """
    Creates a sample text file with mixed content. / 혼합 콘텐츠가 포함된 샘플 텍스트 파일을 생성합니다.
    
    Args:
        tmp_path: Pytest fixture for temporary directory / 임시 디렉토리를 위한 Pytest fixture
        
    Returns:
        Path: Path to the created text file / 생성된 텍스트 파일의 경로
    """
    file_path = tmp_path / "text.txt"
    file_path.write_text("This is a test file.\n한글 테스트", encoding="utf-8")
    return file_path

@pytest.fixture
def sample_pdf_file(tmp_path):
    """
    Creates a sample PDF file with test content. / 테스트 콘텐츠가 포함된 샘플 PDF 파일을 생성합니다.
    
    Args:
        tmp_path: Pytest fixture for temporary directory / 임시 디렉토리를 위한 Pytest fixture
        
    Returns:
        Path: Path to the created PDF file / 생성된 PDF 파일의 경로
        
    Skip:
        If reportlab is not installed / reportlab이 설치되지 않은 경우 스킵
    """
    try:
        from reportlab.pdfgen import canvas
        file_path = tmp_path / "test.pdf"
        
        c = canvas.Canvas(str(file_path))
        c.drawString(100, 750, "This is a test PDF file.")
        c.save()
        
        return file_path
    except ImportError:
        pytest.skip("reportlab is required for PDF tests")

@pytest.fixture
def sample_docx_file(tmp_path):
    """
    Creates a sample DOCX file with text and table content. / 텍스트와 표 콘텐츠가 포함된 샘플 DOCX 파일을 생성합니다.
    
    Args:
        tmp_path: Pytest fixture for temporary directory / 임시 디렉토리를 위한 Pytest fixture
        
    Returns:
        Path: Path to the created DOCX file / 생성된 DOCX 파일의 경로
        
    Skip:
        If python-docx is not installed / python-docx가 설치되지 않은 경우 스킵
    """
    try:
        import docx
        file_path = tmp_path / "test.docx"
        doc = docx.Document()
        
        # Add paragraphs with mixed content / 혼합 콘텐츠로 단락 추가
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("한글 테스트")

        # Create and populate a test table / 테스트 표 생성 및 채우기
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"

        doc.save(file_path)
        return file_path
    except ImportError:
        pytest.skip("python-docx is required for DOCX tests")

def test_get_file_type_text(sample_text_file):
    """Test text file type detection / 텍스트 파일 유형 감지 테스트"""
    file_type = FileHandler.get_file_type(sample_text_file)
    assert file_type == FileType.TEXT

def test_get_file_type_pdf(sample_pdf_file):
    """Test PDF file type detection / PDF 파일 유형 감지 테스트"""
    file_type = FileHandler.get_file_type(sample_pdf_file)
    assert file_type == FileType.PDF
    
def test_get_file_type_docx(sample_docx_file):
    """Test DOCX file type detection / DOCX 파일 유형 감지 테스트"""
    file_type = FileHandler.get_file_type(sample_docx_file)
    assert file_type == FileType.DOCX

def test_get_file_type_unknown(tmp_path):
    """
    Test unknown file type detection with various binary contents. / 다양한 바이너리 콘텐츠로 알 수 없는 파일 유형 감지 테스트
    
    Tests different file signatures and content types: / 다양한 파일 시그니처와 콘텐츠 유형 테스트:
    - PNG signature / PNG 시그니처
    - Random binary / 임의의 바이너리
    - Empty file / 빈 파일
    - ZIP signature / ZIP 시그니처
    """
    cases = [
        (b'\x89PNG\r\n\x1a\n', "PNG signature"),
        (b'Just some random binary \x00\x01\x02', "No magic number"),
        (b'', "Empty file"),
        (b'PK\x03\x04', "ZIP signature")
    ]

    for content, desc in cases:
        unknown_file = tmp_path / f"test_{desc}.xyz"
        unknown_file.write_bytes(content)

        # Log MIME type for debugging / 디버깅을 위한 MIME 타입 로깅
        mime_type = magic.from_file(str(unknown_file), mime=True)
        print(f"\nCase: {desc}")
        print(f"MIME type: {mime_type}")

        file_type = FileHandler.get_file_type(unknown_file)
        assert file_type == FileType.UNKNOWN, f"Failed for {desc}: got {file_type}"

def test_extract_text_from_text_file(sample_text_file):
    """Test text extraction from text file / 텍스트 파일에서 텍스트 추출 테스트"""
    text = FileHandler.extract_text(sample_text_file)
    assert "This is a test file." in text
    assert "한글 테스트" in text

def test_extract_text_from_pdf_file(sample_pdf_file):
    """Test text extraction from PDF file / PDF 파일에서 텍스트 추출 테스트"""
    try:
        import pdfplumber
        text = FileHandler.extract_text(sample_pdf_file)
        assert "This is a test PDF file." in text
    except ImportError:
        pytest.skip("pdfplumber is required for PDF tests")

def test_extract_text_from_docx(sample_docx_file):
    """Test text extraction from DOCX file / DOCX 파일에서 텍스트 추출 테스트"""
    try:
        text = FileHandler.extract_text(sample_docx_file)
        assert "This is a test document." in text
        assert "한글 테스트" in text
        assert "Cell 1" in text and "Cell 4" in text
    except ImportError:
        pytest.skip("python-docx is required for DOCX tests")

def test_extract_text_from_invalid_file(tmp_path):
    """Test text extraction from invalid file / 유효하지 않은 파일에서 텍스트 추출 테스트"""
    invalid_file = tmp_path / "invalid.xyz"
    invalid_file.write_bytes(b'\x89PNG\r\n\x1a\n')  # Use PNG signature / PNG 시그니처 사용
    result = FileHandler.extract_text(invalid_file)
    assert result is None

def test_extract_text_from_nonexistent_file():
    """Test text extraction from nonexistent file / 존재하지 않는 파일에서 텍스트 추출 테스트"""
    try:
        result = FileHandler.extract_text("nonexistent.txt")
        assert result is None
    except FileNotFoundError:
        pass  # Expected behavior / 예상된 동작